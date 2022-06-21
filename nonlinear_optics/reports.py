from datetime import datetime
from typing import List, Optional
from nonlinear_optics import Sample
import matplotlib.pyplot as plt
from docx.document import Document


class Report:

    def __init__(self, samples: List[Sample], doc: Optional[Document] = None):
        self.date = datetime.now()
        self.samples = samples
        self.doc = doc
        if self.doc:
            self.doc.add_heading('Ανάλυση δεδομένων', 0)
        
        self.dots = ['ro', 'bo', 'go', 'yo']

    def create_table(self):
        
        self.doc.add_paragraph('Table 1. Nonlinear optical parameters (Z-scan: 40fs, 800 nm)')
        table = self.doc.add_table(rows=len(self.samples)+1, cols=9)
        
        table.style = 'LightShading-Accent1'
        table.cell(0, 0).text='Samples'
        table.cell(0, 1).text='Concentration'
        table.cell(0, 2).text="γ' (10-21m2/W)"
        table.cell(0, 3).text="Re(x(3)) (10-16 esu)"
        table.cell(0, 4).text="Reγ (10-33 esu)"
        table.cell(0, 5).text="β (&10^{-15}& m/W)"
        table.cell(0, 6).text="Imx(3) (10-16 esu)"
        table.cell(0, 7).text="x(3) (10-16 esu)"
        table.cell(0, 8).text="x(3)/c (10-16 esu/mM)"
        
        for i, sample in enumerate(self.samples):
            table.cell(i+1, 0).text = sample.name
            table.cell(i+1, 1).text = '% 6.2f' % sample.concentration
            table.cell(i+1, 2).text = '% 6.2f' % sample.experiment.gamma
            table.cell(i+1, 3).text = '% 6.2f' % sample.experiment.realXi
            table.cell(i+1, 4).text = '% 6.2f' % sample.experiment.gamma
            table.cell(i+1, 5).text = '% 6.2f' % sample.experiment.beta
            table.cell(i+1, 6).text = '% 6.2f' % sample.experiment.imaginaryXi
            table.cell(i+1, 7).text = '% 6.2f' % sample.experiment.xi
            table.cell(i+1, 8).text = '% 6.2f' % (sample.experiment.xi/sample.concentration)
    
    def plot_spectra(self, show: bool = True, file: Optional[str] = None, documented: bool = False):
        for i, sample in enumerate(self.samples):
            if sample.spectra:
                sample.spectra.plot(show=False, dot=self.dots[i])
        
        if documented:
            self.doc.add_paragraph("Fig. 1. UV vis spectrum of the molecules")
            plt.savefig('./output/images/spectrums.png')
            self.doc.add_picture('./output/images/spectrums.png')

        if show:
            plt.show()
        if file:
            plt.savefig(file, dpi=600)
        
        plt.clf()
        

    def plot_dtpvs(self, show: bool = True, file: Optional[str] = None, documented: bool = False):
    
        for i, sample in enumerate(self.samples):
            sample.experiment.plot_dtpvs(show=False, dot=self.dots[i])
        
        if documented:
            self.doc.add_paragraph('Fig. 2. ΔΤp-v values as a function of the incident laser energy')
            plt.savefig('./output/images/dtpvs.png')
            self.doc.add_picture('./output/images/dtpvs.png')
        if show:
            plt.show()
        if file:
            plt.savefig(file, dpi = 600)
        
        plt.clf()
    
    def add_scans(self):
        for i, sample in enumerate(self.samples):
            for j, scan in enumerate(sample.experiment.zscans):
                scan.plot_scan(show=False)
                plt.savefig(f'./output/images/{i}{j}.png')
                self.doc.add_paragraph(f'sample {sample.name} c={sample.concentration} energy={scan.energy}')
                self.doc.add_picture(f'./output/images/{i}{j}.png')
                plt.clf()